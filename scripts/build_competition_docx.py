from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "competition_submission"
BUILD_DIR = ROOT / "docs" / "competition_submission" / "_build"
OUT_DOCX = DOCS_DIR / "LearnForgeV2_A3_配套文档完整版.docx"
OUT_MD = BUILD_DIR / "LearnForgeV2_A3_配套文档完整版.md"
REFERENCE_DOCX = BUILD_DIR / "competition_reference.docx"

ORDERED_DOCS = [
    "README.md",
    "00_赛题要求对照总览.md",
    "01_需求分析说明书.md",
    "02_系统开发说明书.md",
    "03_测试说明书.md",
    "04_部署运行说明书.md",
    "05_开源与AI工具使用说明.md",
    "06_赛题符合性自查表.md",
    "07_课程知识库与数据说明.md",
    "08_初赛提交材料清单.md",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table) -> None:
    table.style = "Table Grid"
    set_cell_margins(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def build_reference_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 8),
        ("Subtitle", 12, "4B5563", 0, 12),
        ("Heading 1", 16, "1F4D78", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if name != "Subtitle":
            style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Courier New"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        code.font.size = Pt(8.5)
        code.font.color.rgb = RGBColor(17, 24, 39)
        code.paragraph_format.space_before = Pt(3)
        code.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "LearnForge V2 A3 配套文档"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(75, 85, 99)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("第 ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer_run._r.append(fld_begin)
    footer_run._r.append(instr)
    footer_run._r.append(fld_end)
    footer.add_run(" 页")
    for run in footer.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(75, 85, 99)

    sample = doc.add_table(rows=2, cols=3)
    sample.cell(0, 0).text = "项目"
    sample.cell(0, 1).text = "内容"
    sample.cell(0, 2).text = "状态"
    sample.cell(1, 0).text = "样例"
    sample.cell(1, 1).text = "样例内容"
    sample.cell(1, 2).text = "通过"
    style_table(sample)
    doc.save(REFERENCE_DOCX)


def normalize_markdown(text: str, title: str) -> str:
    text = text.replace("\r\n", "\n")
    text = re.sub(r"^# ", "## ", text, count=1, flags=re.MULTILINE)
    text = re.sub(r"^## ", "### ", text, flags=re.MULTILINE)
    text = re.sub(r"^### ", "#### ", text, flags=re.MULTILINE)
    text = text.replace("```text", "```")
    return f"\\newpage\n\n# {title}\n\n{text.strip()}\n"


def build_combined_markdown() -> None:
    parts = [
        "---",
        "title: LearnForge V2 A3 配套文档完整版",
        "subtitle: 基于大模型的个性化资源生成与学习多智能体系统开发",
        "author: LearnForge V2 项目组",
        "date: 2026-06-08",
        "toc: true",
        "toc-depth: 3",
        "---",
        "",
        "# LearnForge V2 A3 配套文档完整版",
        "",
        "赛题：A3-基于大模型的个性化资源生成与学习多智能体系统开发  ",
        "组别：A组（本科、研究生、高职）  ",
        "出题企业：科大讯飞股份有限公司  ",
        "系统名称：LearnForge V2  ",
        "文档日期：2026-06-08",
        "",
        "\\newpage",
        "",
    ]
    for name in ORDERED_DOCS:
        path = DOCS_DIR / name
        title = path.stem
        parts.append(normalize_markdown(path.read_text(encoding="utf-8"), title))
    OUT_MD.write_text("\n".join(parts), encoding="utf-8")


def patch_docx_after_pandoc(path: Path) -> None:
    doc = Document(path)
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Source Code", "Code Block"}:
            paragraph.style = doc.styles["Code Block"]
            for run in paragraph.runs:
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(8.5)
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = "Calibri"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.save(path)


def main() -> None:
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    build_reference_docx()
    build_combined_markdown()
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise SystemExit("pandoc not found")
    cmd = [
        pandoc,
        str(OUT_MD),
        "--from",
        "markdown+pipe_tables+fenced_code_blocks+yaml_metadata_block",
        "--to",
        "docx",
        "--reference-doc",
        str(REFERENCE_DOCX),
        "--toc",
        "--toc-depth=3",
        "-o",
        str(OUT_DOCX),
    ]
    subprocess.run(cmd, check=True)
    patch_docx_after_pandoc(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
